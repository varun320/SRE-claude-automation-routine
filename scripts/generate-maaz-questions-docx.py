"""Generate the Maaz questions DOCX for routine approval walkthrough."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- page margins ----
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ---- base style ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1, color=(20, 30, 50)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        run.font.size = Pt(24)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.font.color.rgb = RGBColor(*color)
    elif level == 1:
        run.font.size = Pt(15)
        run.bold = True
        run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(4)
    elif level == 2:
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(10)
    return p


def para(text, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def question_block(num, total, title, why, options, default, space_lines=4):
    """One question with title, why, options, default, and answer space."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run(f"Question {num} of {total}  ·  {title}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(20, 30, 50)

    p = doc.add_paragraph()
    r = p.add_run("Why we're asking:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(95, 105, 120)
    r2 = p.add_run(why)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(60, 65, 75)

    p = doc.add_paragraph()
    r = p.add_run("Options:")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(95, 105, 120)

    for opt in options:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.8)
        r = p.add_run(opt)
        r.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run("If you don't pick one, we'll default to:  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(95, 105, 120)
    r2 = p.add_run(default)
    r2.italic = True
    r2.font.size = Pt(10.5)

    p = doc.add_paragraph()
    r = p.add_run("Your answer:")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(95, 105, 120)
    for _ in range(space_lines):
        doc.add_paragraph("                                                                                                                                                              ").paragraph_format.space_after = Pt(0)

    # subtle separator
    p = doc.add_paragraph("_" * 80)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.color.rgb = RGBColor(200, 205, 215)


# ============================================================
# TITLE PAGE
# ============================================================
heading("Routine Approval — Questions for Maaz", level=0, color=(15, 25, 55))
para("Phase 1 of the SRE Claude Routines · 12 decisions needed before scheduling", italic=True, size=11.5, color=(95, 105, 120))

# spacer
for _ in range(2):
    doc.add_paragraph("")

heading("What this is", level=1)
para(
    "Before we turn on the first set of automated routines (R1 inbox scan, R2 daily setup, R5 ME inbox scan, R8 Sunday planning), there are 12 small decisions only you can make. They're not technical — they're about how you actually work, what's confidential, and where you want the output to land."
)
para(
    "Each question has 2-4 options. If you don't have a preference, the \"default\" option is what we'll use. Total time: ~20 minutes. Answers go directly to the routine setup."
)

heading("How to answer", level=1)
para("Three options:")
para("   1. Write your answer directly in this Word document and email it back.", size=11)
para("   2. Tell Mohammad in a call — he'll fill it in.", size=11)
para("   3. Reply to each question by number in an email (e.g. \"Q3: option B\").", size=11)

heading("If a question doesn't make sense", level=1)
para("Skip it. Mohammad will follow up on anything left blank. Don't try to figure out a \"right\" answer — your gut preference IS the right answer.")

doc.add_page_break()

# ============================================================
# SECTION 1 — ABOUT YOU AND YOUR TEAM
# ============================================================
heading("Section 1 — About you and your team", level=1, color=(30, 80, 110))
para("Four questions about who's active, what your role is, and how the routines should think about you.", italic=True, color=(95, 105, 120))

question_block(
    num=1, total=12,
    title="Your role title for the routines",
    why="Your Microsoft 365 directory says \"Analyzer Expert.\" The execution plan calls you \"CEO / GM.\" The routines treat you as the decision-maker. Which framing do you want the automated outputs to reflect?",
    options=[
        "A. CEO / General Manager (what the routines + memos assume today)",
        "B. Analyzer Expert (what your Microsoft directory says)",
        "C. Both — different titles for different contexts (tell us when)",
    ],
    default="A — CEO / GM, with the Microsoft directory left as-is for technical work."
)

question_block(
    num=2, total=12,
    title="Inshan — is he still on the team?",
    why="His Microsoft account is marked disabled, but we still see mail flowing into his inbox (a vehicle service report from Toyota Madinat Zayed, for example). The execution plan lists him as the Qatar/Kuwait compliance person, but no Qatar or Kuwait meeting in the last two years involves him. We don't want a routine surfacing his mail if he's no longer active.",
    options=[
        "A. Active — keep the routines treating Inshan as on-staff. We'll fix the Microsoft account.",
        "B. Inactive — stop including his mail in routine outputs.",
        "C. Inactive in some roles but still on payroll — Mohammad to call you to clarify.",
    ],
    default="B — exclude his mail until you say otherwise."
)

question_block(
    num=3, total=12,
    title="Don Green's calendar — why do you have full access?",
    why="Your Microsoft account has read + edit access to Don Green's full calendar. Don is active in our records but not in the original team description we worked from. Just want to confirm what the relationship is so a routine doesn't accidentally book over his calendar or surface his meetings as yours.",
    options=[
        "A. Don is on the team — full access is intentional.",
        "B. Don is a contractor / advisor — the access was granted for a specific purpose (please say what).",
        "C. I didn't know I had this access — please remove it.",
    ],
    default="A — leave the access as-is, treat Don's calendar separate from yours in the routines."
)

question_block(
    num=4, total=12,
    title="Your manager + direct reports in the org chart",
    why="The Microsoft system has no \"who reports to whom\" data populated for SRE. A routine that needs to escalate something (e.g., \"Maaz is on vacation, who handles this?\") has no fallback today. Want us to set up the basic org chart?",
    options=[
        "A. Yes — populate it. Manager = Dwayne Vinck (Board Chairman). Direct reports = Ashley, Dharmesh, Talha, Ron, Chuck, Inshan, plus newer hires.",
        "B. Yes but different mapping — please specify.",
        "C. No — keep the org chart empty.",
    ],
    default="A — set the chart per option A."
)

doc.add_page_break()

# ============================================================
# SECTION 2 — HOW ROUTINES SHOULD BEHAVE
# ============================================================
heading("Section 2 — How the routines should behave", level=1, color=(30, 80, 110))
para("Two questions about the line between \"summarize for you\" and \"take action on your behalf.\"", italic=True, color=(95, 105, 120))

question_block(
    num=5, total=12,
    title="Should routines just summarize, or also start applying rules?",
    why="Your inbox has 24,763 messages, 8,692 of them unread. Zero inbox rules are set up. Routine R1 reads inbound mail and tells you what's important — but we could also let it act: archive obvious junk, move LinkedIn newsletters out of the way, flag customer threads. \"Read-only summarize\" is safe but only treats the symptom. Letting it act treats the cause.",
    options=[
        "A. Summarize only. I'll handle the inbox myself.",
        "B. Summarize + auto-archive obvious junk (LinkedIn newsletters, alarm.com notifications, marketing).",
        "C. Summarize + auto-archive junk + auto-flag known customer threads for follow-up.",
        "D. Full triage including drafting replies (not sending — drafts only).",
    ],
    default="A — summarize only for the first 2 weeks of pilot. If it works, we'll talk about expanding to option B.",
    space_lines=5
)

question_block(
    num=6, total=12,
    title="When the routine reads \"your inbox\" — what counts as yours?",
    why="Because you have full access to Don Green's, Inshan's, and the info@ shared mailbox, all of their incoming mail shows up in your Outlook view. A routine looking at your inbox could include or exclude that — but if it includes it, your morning summary fills up with LinkedIn newsletters meant for Don. Two options:",
    options=[
        "A. Only mail sent directly TO me (maaz@). Drop everything else.",
        "B. My mail + the info@ shared mailbox (since that's company-wide).",
        "C. My mail + info@ + info-me@ (since info-me@ is the ME-region inbox).",
        "D. All of it — I want to see Don's and Inshan's mail too.",
    ],
    default="C — your direct mail, plus info@ and info-me@ for company-wide signals."
)

doc.add_page_break()

# ============================================================
# SECTION 3 — SCOPE
# ============================================================
heading("Section 3 — What's in scope", level=1, color=(30, 80, 110))
para("Two questions about what data the routines pull from.", italic=True, color=(95, 105, 120))

question_block(
    num=7, total=12,
    title="Which shared mailboxes should routines watch?",
    why="SRE has 10 active shared mailboxes (ar@, ap@, info@, info-me@, sales@, careers@, apple@, HusamsBookingPageSRE@, boardroom@, plus a disabled scanner one). Each routine can include or skip them. Suggest including only the ones that map to a real routine job:",
    options=[
        "A. AR (for the Thursday AR routine) + info@ (company-wide inbound) + info-me@ (ME inbound) + sales@ (lead capture) — these are the ones with real signal.",
        "B. All 10 — be thorough.",
        "C. None — only my personal inbox.",
        "D. Different list — please specify.",
    ],
    default="A — those four mailboxes, leave the others alone."
)

question_block(
    num=8, total=12,
    title="Zoho One — connect it to the routines?",
    why="Your Microsoft intranet says explicitly: \"Zoho One is replacing CRM, Timesheets, Project.\" The routines today read Outlook + Calendar + Teams + OneDrive. They DON'T look at Zoho yet — which means anything that lives only in Zoho (customer notes, project status, time sheets, tasks) is invisible to them. Connecting Zoho would make routine outputs much richer, but it's an additional setup step.",
    options=[
        "A. Yes — connect Zoho. The routines should see customer status from Zoho CRM and project status from Zoho Projects.",
        "B. Not yet — keep the routines Microsoft-only for now. Connect Zoho after pilot is stable.",
        "C. Never — Zoho should stay separate from the automation.",
    ],
    default="B — Microsoft-only for pilot. Revisit Zoho at end of Phase 1."
)

doc.add_page_break()

# ============================================================
# SECTION 4 — CALENDAR
# ============================================================
heading("Section 4 — Your calendar and schedule", level=1, color=(30, 80, 110))
para("One important question about how strictly to enforce the Weekly OS schedule.", italic=True, color=(95, 105, 120))

question_block(
    num=9, total=12,
    title="The Day Clock — plan vs reality",
    why="Your Weekly OS plan defines daily standing meetings at 14:05 (Mon = Ashley NA 1:1, Tue = Sales 1:1, Wed = Engineering standup, Thu = Bookkeeper, Fri = Top-5). We looked at your actual calendar for the last 2 years — those exact meetings aren't there. The real pattern is: Thursday 10:00 Bi-Weekly Job meeting (Ashley organizes), a weekly Maaz × Torstein 1:1, and a Monthly Pitstop. Same for calendar tags: the plan says use [Client]/[AIMS]/[Buyer] prefixes. None of your events are tagged. We need to decide which world the routines live in.",
    options=[
        "A. Implement the plan. Routines schedule the missing 14:05 meetings and apply the [Client] tags going forward.",
        "B. Update the plan to match reality. Routines work with the actual cadence you have now (Thu 10:00 Bi-Weekly Job, etc.).",
        "C. Both. Keep the plan as a target; routines work with reality for now, and you'll evolve toward the plan.",
    ],
    default="C — work with reality but keep the plan as the aspiration."
)

# ============================================================
# SECTION 5 — CONFIDENTIAL MATERIAL
# ============================================================
heading("Section 5 — Confidential material", level=1, color=(30, 80, 110))
para("One question about what the routines should never mention out loud.", italic=True, color=(95, 105, 120))

question_block(
    num=10, total=12,
    title="The \"SRE DD\" OneDrive folder — sale process?",
    why="You have a folder at the top of your OneDrive called \"SRE DD\" (23 files, 30 MB, last updated Oct 2025). \"DD\" usually means Due Diligence — typically sale process or board material. We also see ~16 recurring weekly 1:1s in your calendar with someone named Torstein with subjects like \"Finalize DD tracker\" and \"Project P&L and 2026 Pipeline.\" We want to make absolutely sure no routine ever mentions any of this in any output that isn't strictly private to you.",
    options=[
        "A. Confirmed — \"SRE DD\" is sale process material. Routines must never reference it, and must redact anything related (Torstein 1:1 subjects, P&L content, board prep).",
        "B. It's confidential but not sale-process — please clarify so we treat it correctly.",
        "C. It's not actually confidential — routines can mention it.",
    ],
    default="A — treat as fully confidential, never surface in any output."
)

doc.add_page_break()

# ============================================================
# SECTION 6 — OUTPUT
# ============================================================
heading("Section 6 — Where should routine outputs go?", level=1, color=(30, 80, 110))
para("One question — possibly the most important one.", italic=True, color=(95, 105, 120))

question_block(
    num=11, total=12,
    title="Where do you actually want to read the routine output?",
    why="Each routine produces something — a daily inbox summary, a Thursday AR brief, a Sunday week-ahead plan. We need to put that output somewhere you'll actually see it. Options range from \"a file on your computer that you open when you remember\" to \"a Teams chat that pings you the second it's ready.\" Your audit shows Microsoft ToDo is effectively dead (1 task in 12 weeks) and OneNote is dead (last edit Feb 2024). So those aren't real options.",
    options=[
        "A. A Teams chat with myself (or a bot) — I'll see it the moment it lands.",
        "B. An email to myself — it'll show up in my Outlook inbox.",
        "C. A daily/weekly note saved to my OneDrive — I'll open it when I want.",
        "D. A task in Zoho — slots into my existing workflow.",
        "E. A task in ClickUp — slots into project tracking.",
        "F. Different per routine — please specify by routine.",
    ],
    default="A — Teams chat with yourself for daily routines, OneDrive notes for weekly/monthly memos.",
    space_lines=6
)

# ============================================================
# SECTION 7 — SMALL CLARIFICATIONS
# ============================================================
heading("Section 7 — Small clarification", level=1, color=(30, 80, 110))

question_block(
    num=12, total=12,
    title="A small typo to confirm",
    why="We saw a meeting on your calendar from April 20 called \"SRE AI\" with an attendee at the email address info@calldental.ai. Mohammad's email is info@callreceptionist.ai. Just confirming whether calldental.ai is a real other contact, or a typo we should ignore.",
    options=[
        "A. It's a typo — should be info@callreceptionist.ai. Ignore calldental.ai.",
        "B. It's a real address — calldental.ai is a separate venture / contact.",
        "C. I don't recognize either one.",
    ],
    default="A — treat as a typo for info@callreceptionist.ai."
)

# ============================================================
# CLOSING
# ============================================================
doc.add_page_break()
heading("That's all the questions.", level=1, color=(15, 25, 55))
para("Once we have your answers, the routines move from \"draft\" to \"pilot\" — we'll schedule the first one (R1 NA Inbox Scan) and run it daily for 5-10 fires while you review each output. If the output is what you wanted, R1 graduates to \"active\" and we add R2, R5, R8 the same way.")

para("")
para("Total time from your sign-off to first scheduled fire: about 30 minutes (15 minutes connecting Microsoft 365 to the cloud agent system, 15 minutes registering the cron).")

para("")
para("Thanks for the time, Maaz.", italic=True)
para("Mohammad · 2026-05-22", italic=True, color=(95, 105, 120))

# save
out_path = r"D:\projects\prodigy-ai\projects\routines-claude\Maaz-Routine-Approval-Questions.docx"
doc.save(out_path)
print(f"WROTE: {out_path}")
